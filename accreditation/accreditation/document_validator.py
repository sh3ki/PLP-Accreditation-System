"""
Document Validator Utility
Validates uploaded Word documents by comparing their headers with a template document

Enhanced Features:
- Header text validation (case-sensitive)
- Image/logo validation
- Font family validation
- Font size validation
- Text formatting validation (bold, italic, underline)
- Text alignment validation
"""

import os
from docx import Document
from docx.shared import Pt
from typing import Tuple, Optional, Dict, List
from django.conf import settings


class DocumentHeaderValidator:
    """
    Validates Word document headers against a template with strict formatting checks
    """
    
    TEMPLATE_PATH = os.path.join(
        settings.BASE_DIR, 
        'Template.docx'
    )
    
    def __init__(self):
        """Initialize validator and cache template header structure"""
        self._template_header_structure = None
        self._load_template_header()
    
    def _load_template_header(self):
        """Load and cache the template header structure with full formatting"""
        try:
            if not os.path.exists(self.TEMPLATE_PATH):
                raise FileNotFoundError(f"Template file not found at: {self.TEMPLATE_PATH}")
            
            template_doc = Document(self.TEMPLATE_PATH)
            self._template_header_structure = self._extract_header_structure(template_doc)
            
            if not self._template_header_structure['paragraphs'] and not self._template_header_structure['images']:
                raise ValueError("Template document has no header content")
                
        except Exception as e:
            raise Exception(f"Failed to load template header: {str(e)}")
    
    def _extract_header_structure(self, doc: Document) -> Dict:
        """
        Extract complete header structure including text, formatting, and images
        
        Args:
            doc: Document object
            
        Returns:
            Dictionary containing header structure with formatting details
        """
        structure = {
            'paragraphs': [],
            'images': [],
            'tables': []
        }
        
        # Iterate through all sections in the document
        for section in doc.sections:
            header = section.header
            
            # Extract paragraphs with full formatting
            for paragraph in header.paragraphs:
                para_info = self._extract_paragraph_formatting(paragraph)
                if para_info['text'] or para_info['has_image']:
                    structure['paragraphs'].append(para_info)
            
            # Extract images from header
            for rel in header.part.rels.values():
                if "image" in rel.target_ref:
                    structure['images'].append({
                        'rel_id': rel.rId,
                        'target': rel.target_ref
                    })
            
            # Extract tables in header (if any)
            for table in header.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            para_info = self._extract_paragraph_formatting(paragraph)
                            row_data.append(para_info)
                    table_data.append(row_data)
                structure['tables'].append(table_data)
        
        return structure
    
    def _extract_paragraph_formatting(self, paragraph) -> Dict:
        """
        Extract text and formatting details from a paragraph
        
        Args:
            paragraph: Paragraph object from docx
            
        Returns:
            Dictionary with text and formatting information
        """
        para_info = {
            'text': paragraph.text.strip(),
            'alignment': str(paragraph.alignment) if paragraph.alignment else None,
            'runs': [],
            'has_image': False
        }
        
        # Extract run-level formatting (font, size, bold, etc.)
        for run in paragraph.runs:
            # Check if run contains an image
            if 'graphic' in run._element.xml:
                para_info['has_image'] = True
            
            run_info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None,
            }
            para_info['runs'].append(run_info)
        
        return para_info
    
    def _compare_paragraph_formatting(self, template_para: Dict, uploaded_para: Dict) -> Tuple[bool, List[str]]:
        """
        Compare formatting between template and uploaded paragraph
        
        Args:
            template_para: Template paragraph formatting
            uploaded_para: Uploaded paragraph formatting
            
        Returns:
            Tuple of (matches, list of differences)
        """
        differences = []
        
        # Compare text (CASE-SENSITIVE)
        if template_para['text'] != uploaded_para['text']:
            differences.append(f"Text mismatch: Expected '{template_para['text']}' but got '{uploaded_para['text']}'")
        
        # Compare alignment
        if template_para['alignment'] != uploaded_para['alignment']:
            differences.append(f"Alignment mismatch: Expected '{template_para['alignment']}' but got '{uploaded_para['alignment']}'")
        
        # Compare runs (character-level formatting)
        if len(template_para['runs']) != len(uploaded_para['runs']):
            differences.append(f"Run count mismatch: Expected {len(template_para['runs'])} runs but got {len(uploaded_para['runs'])}")
        else:
            for idx, (template_run, uploaded_run) in enumerate(zip(template_para['runs'], uploaded_para['runs'])):
                # Compare text in run (CASE-SENSITIVE)
                if template_run['text'] != uploaded_run['text']:
                    differences.append(f"Run {idx} text mismatch: Expected '{template_run['text']}' but got '{uploaded_run['text']}'")
                
                # Compare bold
                if template_run['bold'] != uploaded_run['bold']:
                    differences.append(f"Run {idx} bold mismatch: Expected {template_run['bold']} but got {uploaded_run['bold']}")
                
                # Compare italic
                if template_run['italic'] != uploaded_run['italic']:
                    differences.append(f"Run {idx} italic mismatch: Expected {template_run['italic']} but got {uploaded_run['italic']}")
                
                # Compare underline
                if template_run['underline'] != uploaded_run['underline']:
                    differences.append(f"Run {idx} underline mismatch: Expected {template_run['underline']} but got {uploaded_run['underline']}")
                
                # Compare font name
                if template_run['font_name'] != uploaded_run['font_name']:
                    differences.append(f"Run {idx} font mismatch: Expected '{template_run['font_name']}' but got '{uploaded_run['font_name']}'")
                
                # Compare font size (with tolerance of 0.1pt for floating point issues)
                if template_run['font_size'] and uploaded_run['font_size']:
                    if abs(template_run['font_size'] - uploaded_run['font_size']) > 0.1:
                        differences.append(f"Run {idx} font size mismatch: Expected {template_run['font_size']}pt but got {uploaded_run['font_size']}pt")
                elif template_run['font_size'] != uploaded_run['font_size']:
                    differences.append(f"Run {idx} font size mismatch: Expected {template_run['font_size']} but got {uploaded_run['font_size']}")
        
        return (len(differences) == 0, differences)
    
    def validate_document(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate if the uploaded document's header matches the template exactly
        including text, formatting, fonts, sizes, and images
        
        Args:
            file_path: Path to the uploaded document file
            
        Returns:
            Tuple of (is_valid, error_message)
            - is_valid: True if header matches exactly, False otherwise
            - error_message: Empty string if valid, detailed differences if invalid
        """
        try:
            # Load uploaded document
            uploaded_doc = Document(file_path)
            uploaded_structure = self._extract_header_structure(uploaded_doc)
            
            if not uploaded_structure['paragraphs'] and not uploaded_structure['images']:
                return False, "Uploaded document has no header content"
            
            all_differences = []
            
            # Compare number of paragraphs
            template_paras = self._template_header_structure['paragraphs']
            uploaded_paras = uploaded_structure['paragraphs']
            
            if len(template_paras) != len(uploaded_paras):
                return False, (
                    f"Header structure mismatch: Template has {len(template_paras)} paragraphs "
                    f"but uploaded document has {len(uploaded_paras)} paragraphs. "
                    "Please use the exact Template.docx header."
                )
            
            # Compare each paragraph with formatting
            for idx, (template_para, uploaded_para) in enumerate(zip(template_paras, uploaded_paras)):
                matches, differences = self._compare_paragraph_formatting(template_para, uploaded_para)
                if not matches:
                    all_differences.append(f"Paragraph {idx + 1}:")
                    all_differences.extend([f"  - {diff}" for diff in differences])
            
            # Compare images
            template_images = self._template_header_structure['images']
            uploaded_images = uploaded_structure['images']
            
            if len(template_images) != len(uploaded_images):
                all_differences.append(
                    f"Image count mismatch: Template has {len(template_images)} images "
                    f"but uploaded document has {len(uploaded_images)} images"
                )
            
            # Compare tables
            template_tables = self._template_header_structure['tables']
            uploaded_tables = uploaded_structure['tables']
            
            if len(template_tables) != len(uploaded_tables):
                all_differences.append(
                    f"Table count mismatch: Template has {len(template_tables)} tables "
                    f"but uploaded document has {len(uploaded_tables)} tables"
                )
            
            # If there are any differences, return detailed error
            if all_differences:
                error_message = (
                    "Document header does not match the required template exactly. "
                    "Please ensure you copy the header from Template.docx exactly (including fonts, sizes, and formatting).\n\n"
                    "Differences found:\n" + "\n".join(all_differences)
                )
                return False, error_message
            
            return True, ""
                
        except Exception as e:
            return False, f"Error validating document: {str(e)}"
    
    def validate_uploaded_file(self, uploaded_file) -> Tuple[bool, str]:
        """
        Validate an uploaded Django file object
        
        Args:
            uploaded_file: Django UploadedFile object
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        import tempfile
        
        try:
            # Create a temporary file to save the upload
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                # Write uploaded file content to temp file
                for chunk in uploaded_file.chunks():
                    tmp_file.write(chunk)
                tmp_file_path = tmp_file.name
            
            # Validate the temporary file
            is_valid, error_message = self.validate_document(tmp_file_path)
            
            # Clean up temporary file
            try:
                os.unlink(tmp_file_path)
            except:
                pass
            
            return is_valid, error_message
            
        except Exception as e:
            return False, f"Error processing uploaded file: {str(e)}"
    
    def get_template_header_preview(self) -> str:
        """
        Get a preview of the template header for display/debugging
        
        Returns:
            Formatted string showing template header structure
        """
        if not self._template_header_structure:
            return "No template header loaded"
        
        preview_lines = []
        preview_lines.append("=" * 60)
        preview_lines.append("TEMPLATE HEADER STRUCTURE")
        preview_lines.append("=" * 60)
        
        # Show paragraphs with formatting
        for idx, para in enumerate(self._template_header_structure['paragraphs']):
            preview_lines.append(f"\nParagraph {idx + 1}:")
            preview_lines.append(f"  Text: {para['text']}")
            preview_lines.append(f"  Alignment: {para['alignment']}")
            
            if para['runs']:
                preview_lines.append(f"  Runs: {len(para['runs'])}")
                for run_idx, run in enumerate(para['runs']):
                    if run['text']:
                        formatting = []
                        if run['bold']:
                            formatting.append("Bold")
                        if run['italic']:
                            formatting.append("Italic")
                        if run['underline']:
                            formatting.append("Underline")
                        
                        preview_lines.append(
                            f"    Run {run_idx + 1}: '{run['text']}' | "
                            f"Font: {run['font_name']} | "
                            f"Size: {run['font_size']}pt | "
                            f"Style: {', '.join(formatting) if formatting else 'Normal'}"
                        )
        
        # Show images
        if self._template_header_structure['images']:
            preview_lines.append(f"\nImages: {len(self._template_header_structure['images'])}")
            for idx, img in enumerate(self._template_header_structure['images']):
                preview_lines.append(f"  Image {idx + 1}: {img['target']}")
        
        # Show tables
        if self._template_header_structure['tables']:
            preview_lines.append(f"\nTables: {len(self._template_header_structure['tables'])}")
        
        preview_lines.append("=" * 60)
        
        return "\n".join(preview_lines)


# Global validator instance (singleton pattern)
_validator_instance = None


def get_validator() -> DocumentHeaderValidator:
    """
    Get or create the global validator instance
    
    Returns:
        DocumentHeaderValidator instance
    """
    global _validator_instance
    if _validator_instance is None:
        _validator_instance = DocumentHeaderValidator()
    return _validator_instance


def validate_document_header(uploaded_file) -> Tuple[bool, str]:
    """
    Convenience function to validate an uploaded document's header
    
    Args:
        uploaded_file: Django UploadedFile object
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    validator = get_validator()
    return validator.validate_uploaded_file(uploaded_file)
