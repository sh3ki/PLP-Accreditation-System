"""
Utility functions for adding footer watermarks to Word documents
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io


def get_accreditation_type_short_code(type_name):
    """
    Get short code for accreditation type
    
    Args:
        type_name: Full accreditation type name
        
    Returns:
        str: Short code (ALCUCOA or COPC)
    """
    type_name_lower = type_name.lower()
    
    if 'association of local colleges and university commission on accreditation' in type_name_lower:
        return 'ALCUCOA'
    elif 'certificate of program compliance' in type_name_lower:
        return 'COPC'
    else:
        # Default to first letters of each word if not recognized
        words = type_name.split()
        return ''.join([word[0].upper() for word in words if word])


def get_checklist_number_from_name(checklist_name):
    """
    Extract checklist number from name like "Checklist 1", "Checklist 2", etc.
    
    Args:
        checklist_name: Name of the checklist
        
    Returns:
        str: Checklist number
    """
    try:
        if 'Checklist' in checklist_name or 'checklist' in checklist_name:
            # Extract number from "Checklist X" format
            import re
            match = re.search(r'\d+', checklist_name)
            if match:
                return match.group()
        return '1'  # Default to 1 if not found
    except:
        return '1'


def add_footer_code_to_document(file, metadata, original_filename):
    """
    Add footer code to Word document before uploading to Cloudinary.
    This is for REQUIRED documents only.
    
    Args:
        file: Django UploadedFile object (Word document)
        metadata: Dictionary containing:
            - accreditation_type_name: Full accreditation type name
            - dept_code: Department code (e.g., 'CCS')
            - program_code: Program code (e.g., 'BSCS')
            - area_name: Area name (e.g., 'Area 1')
            - checklist_name: Checklist name (e.g., 'Checklist 2')
            - checklist_id: Checklist ID for version tracking
            - document_name: Original document name for version tracking
        original_filename: Original filename to preserve
            
    Returns:
        tuple: (io.BytesIO object with modified document, new filename with .docx extension)
    """
    try:
        # Load the Word document
        doc = Document(file)
        
        # Generate footer code
        type_short = get_accreditation_type_short_code(metadata.get('accreditation_type_name', ''))
        dept_code = metadata.get('dept_code', 'DEPT')
        program_code = metadata.get('program_code', 'PROG')
        area_name = metadata.get('area_name', 'AREA1').replace(' ', '').upper()  # Remove spaces and uppercase
        checklist_num = get_checklist_number_from_name(metadata.get('checklist_name', 'Checklist 1'))
        
        # Get document number and version
        checklist_id = metadata.get('checklist_id')
        document_name = metadata.get('document_name', original_filename)
        doc_number, version_string = get_document_number_and_version(checklist_id, document_name)
        
        # Format timestamp: MM:DD:YYYY-HH:MM:SS
        now = datetime.now()
        timestamp = now.strftime('%m:%d:%Y-%H:%M:%S')
        
        # Construct the footer code
        # Format: {TYPE_SHORT}-{DEPT_CODE}-{PROG_CODE}-{AREA_NAME}-{CHECKLIST_NUM}-{DOC_NUM}-1v.0.0 {TIMESTAMP}
        footer_code = f"{type_short}-{dept_code}-{program_code}-{area_name}-{checklist_num}-{doc_number}-{version_string} {timestamp}"
        
        # Add footer to all sections
        for section in doc.sections:
            # Get or create footer
            footer = section.footer
            
            # Clear existing footer content (optional, comment out if you want to keep existing content)
            # for paragraph in footer.paragraphs:
            #     paragraph.clear()
            
            # Add new paragraph to footer
            footer_paragraph = footer.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # LEFT align
            
            # Add the footer code text
            run = footer_paragraph.add_run(footer_code)
            run.font.size = Pt(9)  # Small font size
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
            run.font.name = 'Arial'
        
        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Generate new filename with .docx extension
        # Remove original extension and add .docx
        base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
        new_filename = f"{base_name}.docx"
        
        return output, new_filename
        
    except Exception as e:
        print(f"Error adding footer to document: {str(e)}")
        # If error, return original file with original name
        file.seek(0)
        return file, original_filename


def count_existing_required_documents(dept_id, prog_id, type_id, area_id, checklist_id):
    """
    Count existing required documents in a checklist to determine the next document number.
    
    Args:
        dept_id: Department ID
        prog_id: Program ID
        type_id: Accreditation type ID
        area_id: Area ID
        checklist_id: Checklist ID
        
    Returns:
        int: Next document number (1-based)
    """
    try:
        from accreditation.firebase_utils import firestore_helper
        
        # Query documents collection
        docs_ref = firestore_helper.db.collection('documents')
        
        # Filter by checklist and required status
        query = docs_ref.where('checklist_id', '==', checklist_id)\
                       .where('is_required', '==', True)\
                       .where('is_active', '==', True)
        
        # Count existing documents
        existing_docs = list(query.stream())
        
        # Return next number (existing count + 1)
        return len(existing_docs) + 1
        
    except Exception as e:
        print(f"Error counting documents: {str(e)}")
        return 1  # Default to 1 if error


def get_document_number_and_version(checklist_id, document_name):
    """
    Get document number and version for footer based on REQUIRED documents in checklist.
    
    Logic:
    - Count ONLY required documents (is_required=True)
    - Each new required document upload gets the next sequential number
    - Version is always 1v.0.0 for all documents
    
    Args:
        checklist_id: Checklist ID
        document_name: Original document name
        
    Returns:
        tuple: (document_number, version_string)
        
    Example:
        First required doc upload -> (1, "1v.0.0")
        Second required doc upload (even same name) -> (2, "1v.0.0")
        Third required doc upload -> (3, "1v.0.0")
    """
    try:
        from accreditation.firebase_utils import get_all_documents
        
        print(f"[DOC NUMBER DEBUG] Getting doc number for: checklist_id={checklist_id}, document_name={document_name}")
        
        # Get ONLY required documents in this checklist
        all_documents = get_all_documents('documents')
        checklist_docs = [
            d for d in all_documents 
            if d.get('checklist_id') == checklist_id 
            and d.get('is_required', False) == True
            and d.get('is_active', True)
            and not d.get('is_archived', False)
        ]
        
        print(f"[DOC NUMBER DEBUG] Found {len(checklist_docs)} required documents in checklist")
        
        # The next document number is simply the count + 1
        doc_number = len(checklist_docs) + 1
        version_string = "1v.0.0"
        
        print(f"[DOC NUMBER DEBUG] Next required document number: {doc_number}, version: {version_string}")
        return doc_number, version_string
        
    except Exception as e:
        print(f"Error getting document number: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1, "1v.0.0"  # Default to first document
